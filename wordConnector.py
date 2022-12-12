from telnetlib import WONT
from docx import Document
import subprocess
class WordConnector:
    def __init__(self):
        x = "blub"
        self.valueDic = {
            "__date__":x,
            "__Client_Name__":x,
            "__Client_Street_Nr__":x,
            "__Client_Place__":x,
            "__Bill_Nr__":x,
            "__Job__":x,
            "__Time_of_Job__":x,
            "__Sum_Price__":x,
        }

    def editBill(self, docPath, valueDic, wage, positions):
        document = Document(docPath)

        for p in document.paragraphs:
            if "position_1" in p.text:
                positionParagraph = p
                break
        
        moneySum = 0
        defaultText = positionParagraph.text
        i = 0
        while i < len(positions):
            positionParagraph
            title = positions[i][0]
            time_s = positions[i][1]
            time_h_round = round(time_s/3600)
            if time_h_round > 0:
                i+=1
                money_round = time_h_round * wage

                moneySum += money_round

                nTabs = (6-(len(title)//6))
                text = defaultText.replace("__position_1__", title+nTabs*"\t")
                text = text.replace("1.", f"{i}.")
                text = text.replace("__1_N__", self.billify(time_h_round))
                text = text.replace("__1_P__", self.billify(wage))
                text = text.replace("__1_T__", self.billify(money_round))
                print(text, positionParagraph.text, type(document.paragraphs))
                positionParagraph.insert_paragraph_before(text)
            else:
                positions.pop(i)

        # remove default position
        positionParagraph.text = ""

        print(len(document.paragraphs) ,"len3")
        #document.paragraphs.remove(positionParagraph)
        valueDic["__Sum_Price__"] = self.billify(moneySum)

        for p in document.paragraphs:
            for key in self.valueDic:
                if len(p.runs) > 1: 
                    print()
                if key in p.text:
                    collect = ""
                    for r in p.runs:
                        collect += r.text
                        r.text = ""
                    p.runs[-1].text = collect.replace(key, str(valueDic[key]))
                    #p.text = p.text.replace(key, valueDic[key])
        
        savePath = docPath[:-5]+"_copy.docx"
        document.save(savePath)
        subprocess.call(('open', savePath))

    def billify(self, number):
        number = str(number).replace(".", ",")
        number = number if "," in number else f"{number},00"
        while not number[-3] == ",":
            number = f"{number}0"
        return number
